"""文档业务逻辑：安全上传 + 多格式解析（PDF/docx/文本）+ RAG 入库。"""  # 模块级文档字符串，描述文档业务逻辑

import io  # 导入io标准库，用于字节流操作
import os  # 导入操作系统接口标准库

from app.core.config import get_settings  # 导入配置获取函数
from app.core.constants import ALLOWED_DOC_EXTENSIONS, ALLOWED_IMAGE_EXTENSIONS  # 导入允许的文档和图片扩展名集合
from app.core.exceptions import BadRequestError, ConflictError, PayloadTooLargeError  # 导入多种异常类
from app.core.logging import setup_logger  # 导入日志记录器配置函数
from app.core.security import is_allowed_extension, sanitize_filename, validate_upload_path  # 导入安全相关函数
from app.agents.runtime import get_vector_store  # 导入向量库获取函数
from app.memory.rag import semantic_chunk  # 导入语义切片函数

logger = setup_logger("service.document")  # 创建名为service.document的日志记录器


# ============================================================  # 分隔注释
# 安全上传  # 说明该部分为安全上传逻辑
# ============================================================  # 分隔注释

def save_upload(file_bytes: bytes, filename: str, allowed_set: set) -> str:  # 定义安全上传函数
    """
    安全保存上传文件，返回存储后的文件名。

    校验链：扩展名白名单 → 大小限制 → 文件名消毒 → 路径遍历防护。
    """  # 函数文档字符串
    settings = get_settings()  # 获取配置

    if not is_allowed_extension(filename, allowed_set):  # 如果扩展名不在白名单
        raise BadRequestError(f"不支持的文件类型：{filename}")  # 抛出错误请求异常

    if len(file_bytes) > settings.max_upload_bytes:  # 如果文件超过大小限制
        raise PayloadTooLargeError(  # 抛出负载过大异常
            f"文件超过 {settings.MAX_UPLOAD_SIZE_MB}MB 限制"  # 错误消息
        )

    safe_name = sanitize_filename(filename)  # 对文件名进行消毒处理
    filepath = os.path.join(settings.UPLOAD_FOLDER, safe_name)  # 拼接完整文件路径
    if not validate_upload_path(filepath, settings.UPLOAD_FOLDER):  # 校验路径合法性（防遍历）
        raise BadRequestError("非法的文件路径")  # 抛出错误请求异常

    with open(filepath, "wb") as f:  # 以二进制写模式打开文件
        f.write(file_bytes)  # 写入文件内容
    logger.info("文件已保存: %s (%d 字节)", safe_name, len(file_bytes))  # 记录保存日志
    return safe_name  # 返回安全文件名


def save_image_upload(file_bytes: bytes, filename: str) -> str:  # 定义图片上传保存函数
    """保存聊天图片上传。"""  # 函数文档字符串
    return save_upload(file_bytes, filename, ALLOWED_IMAGE_EXTENSIONS)  # 调用通用上传函数，使用图片扩展名集合


# ============================================================  # 分隔注释
# 多格式文档解析  # 说明该部分为多格式文档解析逻辑
# ============================================================  # 分隔注释

def parse_document(file_bytes: bytes, filename: str) -> str:  # 定义统一文档解析函数
    """
    统一文档解析接口：parse_document(bytes, filename) -> str。

    支持 PDF（pdfplumber）、Word（python-docx）及各类纯文本格式。
    """  # 函数文档字符串
    ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""  # 提取文件扩展名并转小写

    if ext == "pdf":  # 如果是PDF
        return _parse_pdf(file_bytes)  # 调用PDF解析函数
    if ext == "docx":  # 如果是Word文档
        return _parse_docx(file_bytes)  # 调用Word解析函数
    return _parse_text(file_bytes, filename)  # 其他类型调用纯文本解析函数


def _parse_pdf(file_bytes: bytes) -> str:  # 定义PDF解析内部函数
    """使用 pdfplumber 提取 PDF 文本（含表格）。"""  # 函数文档字符串
    import pdfplumber  # 延迟导入pdfplumber库

    parts = []  # 初始化文本片段列表
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:  # 打开PDF字节流
        for page_num, page in enumerate(pdf.pages, 1):  # 遍历每一页
            text = page.extract_text() or ""  # 提取页面文本
            if text.strip():  # 如果文本非空
                parts.append(text)  # 添加到片段列表
            # 提取表格（转为简单文本行）  # 内部注释说明表格提取
            for table in page.extract_tables():  # 遍历页面表格
                for row in table:  # 遍历表格行
                    cells = [str(cell).strip() if cell else "" for cell in row]  # 处理每个单元格
                    if any(cells):  # 如果行有非空单元格
                        parts.append(" | ".join(cells))  # 用竖线连接单元格
    if not parts:  # 如果未提取到任何文本
        raise BadRequestError("PDF 未提取到任何文本（可能是扫描件）")  # 抛出异常
    return "\n\n".join(parts)  # 用双换行连接所有片段


def _parse_docx(file_bytes: bytes) -> str:  # 定义Word文档解析内部函数
    """使用 python-docx 提取 Word 段落与表格。"""  # 函数文档字符串
    from docx import Document  # 延迟导入python-docx库

    doc = Document(io.BytesIO(file_bytes))  # 加载Word文档
    parts = []  # 初始化文本片段列表
    for para in doc.paragraphs:  # 遍历段落
        if para.text.strip():  # 如果段落文本非空
            parts.append(para.text)  # 添加到片段列表
    for table in doc.tables:  # 遍历表格
        for row in table.rows:  # 遍历表格行
            cells = [cell.text.strip() for cell in row.cells]  # 处理每个单元格
            if any(cells):  # 如果行有非空单元格
                parts.append(" | ".join(cells))  # 用竖线连接单元格
    if not parts:  # 如果未提取到任何文本
        raise BadRequestError("Word 文档未提取到任何文本")  # 抛出异常
    return "\n\n".join(parts)  # 用双换行连接所有片段


def _parse_text(file_bytes: bytes, filename: str) -> str:  # 定义纯文本解析内部函数
    """纯文本格式解码（txt/md/csv/json/html/py 等）。"""  # 函数文档字符串
    for encoding in ("utf-8", "gbk", "latin-1"):  # 依次尝试多种编码
        try:  # 尝试解码
            text = file_bytes.decode(encoding)  # 用当前编码解码字节
            if text.strip():  # 如果解码结果非空
                return text  # 返回解码文本
        except (UnicodeDecodeError, ValueError):  # 如果解码失败
            continue  # 继续尝试下一个编码
    raise BadRequestError(f"无法解析文件内容：{filename}")  # 所有编码都失败，抛出异常


# ============================================================  # 分隔注释
# RAG 入库与管理  # 说明该部分为RAG入库与管理逻辑
# ============================================================  # 分隔注释

async def ingest_document(file_bytes: bytes, filename: str) -> int:  # 定义文档入库协程函数
    """解析文档 → 语义切片 → 存入 RAG 向量库，返回切片数。"""  # 函数文档字符串
    if not is_allowed_extension(filename, ALLOWED_DOC_EXTENSIONS):  # 如果扩展名不在文档白名单
        raise BadRequestError(f"不支持的文档类型：{filename}")  # 抛出错误请求异常

    text = parse_document(file_bytes, filename)  # 解析文档为文本
    chunks = semantic_chunk(text)  # 对文本进行语义切片
    if not chunks:  # 如果切片结果为空
        raise BadRequestError("文档切片为空，请检查文档内容")  # 抛出异常

    store = get_vector_store()  # 获取向量库实例
    count = await store.add_document_chunks(chunks, filename)  # 将切片添加到向量库
    if count == 0:  # 如果入库失败
        raise ConflictError("向量库暂不可用（嵌入模型未就绪），文档未入库")  # 抛出冲突异常
    logger.info("文档已入库 RAG: %s (%d 个切片)", filename, count)  # 记录入库日志
    return count  # 返回切片数量


async def list_documents() -> list[dict]:  # 定义列出文档协程函数
    """列出所有 RAG 文档（文件名 + 切片数 + 时间）。"""  # 函数文档字符串
    store = get_vector_store()  # 获取向量库实例
    return await store.list_documents()  # 调用向量库列出文档


async def delete_document(filename: str) -> bool:  # 定义删除文档协程函数
    """删除指定 RAG 文档的全部切片。"""  # 函数文档字符串
    store = get_vector_store()  # 获取向量库实例
    return await store.delete_document(filename)  # 调用向量库删除文档
